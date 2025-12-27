# import re

def parse_resume_text(text):
    text = text.lower()

    SKILL_KEYWORDS = [
        "python", "java", "c++", "javascript", "react",
        "sql", "mongodb", "flask", "django", "node"
    ]

    skills = []
    for skill in SKILL_KEYWORDS:
        if skill in text:
            skills.append(skill.capitalize())

    education_found = "education" in text
    experience_found = "experience" in text or "internship" in text

    # 🚨 Resume validation rule
    is_valid_resume = education_found and len(skills) >= 2

    return {
        "skills": list(set(skills)),
        "education_found": education_found,
        "experience_found": experience_found,
        "is_valid_resume": is_valid_resume
    }
# backend/utils/resume_parser.py




# backend/utils/resume_parser.py
# import os
# import pdfplumber
# import docx
# import PyPDF2
# from typing import Optional, Dict
# import re

# class ResumeParser:
#     def __init__(self):
#         self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
    
#     def extract_text(self, file_path: str) -> Optional[str]:
#         """
#         Extract text from resume file
#         """
#         if not os.path.exists(file_path):
#             raise FileNotFoundError(f"File not found: {file_path}")
        
#         _, ext = os.path.splitext(file_path)
#         ext = ext.lower()
        
#         try:
#             if ext == '.pdf':
#                 return self._extract_from_pdf(file_path)
#             elif ext == '.docx':
#                 return self._extract_from_docx(file_path)
#             elif ext == '.txt':
#                 return self._extract_from_txt(file_path)
#             elif ext == '.doc':
#                 # .doc files - try to read as binary text
#                 return self._extract_from_doc(file_path)
#             else:
#                 raise ValueError(f"Unsupported file format: {ext}")
#         except Exception as e:
#             print(f"Error extracting text: {str(e)}")
#             return self._fallback_extraction(file_path)
    
#     def _extract_from_pdf(self, file_path: str) -> str:
#         """Extract text from PDF using pdfplumber"""
#         text = ""
#         try:
#             with pdfplumber.open(file_path) as pdf:
#                 for page in pdf.pages:
#                     page_text = page.extract_text()
#                     if page_text:
#                         text += page_text + "\n"
            
#             # If no text extracted, try PyPDF2 as fallback
#             if not text.strip():
#                 text = self._extract_with_pypdf2(file_path)
                
#         except Exception as e:
#             print(f"pdfplumber extraction failed: {str(e)}")
#             text = self._extract_with_pypdf2(file_path)
        
#         return text
    
#     def _extract_with_pypdf2(self, file_path: str) -> str:
#         """Extract text from PDF using PyPDF2 (fallback)"""
#         text = ""
#         try:
#             with open(file_path, 'rb') as file:
#                 pdf_reader = PyPDF2.PdfReader(file)
#                 for page in pdf_reader.pages:
#                     page_text = page.extract_text()
#                     if page_text:
#                         text += page_text + "\n"
#         except Exception as e:
#             print(f"PyPDF2 extraction failed: {str(e)}")
        
#         return text
    
#     def _extract_from_docx(self, file_path: str) -> str:
#         """Extract text from DOCX"""
#         try:
#             doc = docx.Document(file_path)
#             text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
#             return text
#         except Exception as e:
#             print(f"DOCX extraction error: {str(e)}")
#             return ""
    
#     def _extract_from_doc(self, file_path: str) -> str:
#         """Basic extraction from .doc files"""
#         try:
#             # Try to read as binary and extract text
#             with open(file_path, 'rb') as f:
#                 content = f.read()
#                 # Try to decode as text
#                 try:
#                     return content.decode('utf-8', errors='ignore')
#                 except:
#                     return content.decode('latin-1', errors='ignore')
#         except Exception as e:
#             print(f"DOC extraction error: {str(e)}")
#             return ""
    
#     def _extract_from_txt(self, file_path: str) -> str:
#         """Extract text from TXT"""
#         try:
#             with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
#                 return file.read()
#         except:
#             try:
#                 with open(file_path, 'r', encoding='latin-1', errors='ignore') as file:
#                     return file.read()
#             except Exception as e:
#                 print(f"TXT extraction error: {str(e)}")
#                 return ""
    
#     def _fallback_extraction(self, file_path: str) -> str:
#         """Final fallback method"""
#         try:
#             with open(file_path, 'rb') as f:
#                 content = f.read()
#                 return content.decode('utf-8', errors='ignore')
#         except:
#             return ""
    
#     def extract_sections(self, text: str) -> Dict[str, str]:
#         """
#         Extract common resume sections
#         """
#         sections = {
#             'summary': '',
#             'education': '',
#             'experience': '',
#             'skills': '',
#             'projects': '',
#             'certifications': ''
#         }
        
#         if not text:
#             return sections
        
#         # Simple regex-based section extraction
#         patterns = {
#             'education': r'(?i)(?:education|academic background|qualifications?)[\s:]*\n(.*?)(?=\n\s*\n|\n\s*(?:experience|skills|projects|certifications|$))',
#             'experience': r'(?i)(?:work\s*experience|professional\s*experience|employment|experience)[\s:]*\n(.*?)(?=\n\s*\n|\n\s*(?:education|skills|projects|certifications|$))',
#             'skills': r'(?i)(?:technical\s*skills|skills|competencies|expertise)[\s:]*\n(.*?)(?=\n\s*\n|\n\s*(?:education|experience|projects|certifications|$))',
#             'projects': r'(?i)(?:projects|personal\s*projects|portfolio)[\s:]*\n(.*?)(?=\n\s*\n|\n\s*(?:education|experience|skills|certifications|$))',
#             'certifications': r'(?i)(?:certifications|certificate|licenses|trainings?)[\s:]*\n(.*?)(?=\n\s*\n|\n\s*(?:education|experience|skills|projects|$))'
#         }
        
#         for section, pattern in patterns.items():
#             match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
#             if match:
#                 sections[section] = match.group(1).strip()
        
#         # First 100 chars as summary if no specific summary section
#         sections['summary'] = text[:200].strip() if not sections['summary'] else sections['summary']
        
#         return sections





# utils/resume_parser.py - Simplified without pandas
import os
import logging
import traceback

logger = logging.getLogger(__name__)

def extract_text_from_pdf(pdf_path):
    """Extract text from PDF file"""
    try:
        # Try to import PyPDF2
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            logger.warning("PyPDF2 not installed. Install with: pip install PyPDF2")
            return "PDF text extraction requires PyPDF2. Please install it."
        
        text = ""
        with open(pdf_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        
        logger.info(f"Extracted {len(text)} characters from PDF")
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        return f"Error extracting PDF: {str(e)}"

def extract_text_from_docx(docx_path):
    """Extract text from DOCX file"""
    try:
        # Try to import python-docx
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not installed. Install with: pip install python-docx")
            return "DOCX text extraction requires python-docx. Please install it."
        
        doc = Document(docx_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        logger.info(f"Extracted {len(text)} characters from DOCX")
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        return f"Error extracting DOCX: {str(e)}"

def extract_text_from_txt(txt_path):
    """Extract text from TXT file"""
    try:
        with open(txt_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        logger.info(f"Extracted {len(text)} characters from TXT")
        return text
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {str(e)}")
        return f"Error extracting TXT: {str(e)}"

def extract_text_from_file(file_path):
    """
    Main function to extract text from supported file types
    """
    try:
        if not os.path.exists(file_path):
            return f"File not found: {file_path}"
        
        file_extension = file_path.lower().split('.')[-1]
        
        logger.info(f"Extracting text from {file_path} (type: {file_extension})")
        
        if file_extension == 'pdf':
            return extract_text_from_pdf(file_path)
        elif file_extension in ['docx', 'doc']:
            return extract_text_from_docx(file_path)
        elif file_extension == 'txt':
            return extract_text_from_txt(file_path)
        else:
            return f"Unsupported file type: {file_extension}"
            
    except Exception as e:
        logger.error(f"Failed to extract text: {str(e)}")
        return f"Extraction error: {str(e)}"